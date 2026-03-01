# cv_extractor.py
# Extracts raw text from uploaded PDF or DOCX files.
# This extracted text is fed directly into the Gemini prompt so it
# uses REAL data from the user's CV instead of hallucinating.

import pdfplumber
import io
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts all text from a PDF file.
    pdfplumber is more accurate than PyPDF2 for German text with umlauts.

    file_bytes: raw bytes of the uploaded file (from Streamlit's uploader)
    returns: full text as a single string
    """
    text_chunks = []

    # BytesIO wraps the bytes so pdfplumber can treat it like an open file
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_number, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)

    full_text = "\n".join(text_chunks)

    if not full_text.strip():
        return "⚠️ Could not extract text from PDF. It may be a scanned image."

    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts all text from a DOCX file paragraph by paragraph.
    Also extracts text from tables (for two-column CV layouts).

    file_bytes: raw bytes of the uploaded file
    returns: full text as a single string
    """
    doc = Document(io.BytesIO(file_bytes))
    text_chunks = []

    # Extract regular paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_chunks.append(para.text)

    # Extract text from tables (important: two-column CVs store data in tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_chunks.append(" | ".join(row_text))

    full_text = "\n".join(text_chunks)

    if not full_text.strip():
        return "⚠️ Could not extract text from DOCX file."

    return full_text


def extract_cv_text(uploaded_file) -> str:
    """
    Main entry point — detects file type and routes to correct extractor.

    uploaded_file: Streamlit UploadedFile object
    returns: extracted text string
    """
    if uploaded_file is None:
        return None

    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)

    elif filename.endswith(".docx") or filename.endswith(".doc"):
        return extract_text_from_docx(file_bytes)

    else:
        return f"⚠️ Unsupported file format: {uploaded_file.name}"
