# document_builder.py
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


def create_cv_document(cv_text: str, full_name: str, photo_path: str = None) -> io.BytesIO:
    """
    Creates a two-column German CV with photo (if provided).
    Tight spacing (6pt between sections, no extra line breaks).
    """
    if not cv_text or not isinstance(cv_text, str):
        raise ValueError(f"cv_text must be a non-empty string, got: {type(cv_text)}")

    doc = Document()

    # DIN 5008 margins (2.5cm all around except 2cm right)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)

    # --- HEADER WITH NAME + PHOTO ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False

    # Left cell: Name
    left_cell = header_table.rows[0].cells[0]
    left_cell.width = Inches(4.5)
    name_para = left_cell.paragraphs[0]
    name_run = name_para.add_run(full_name.upper())
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.name = "Arial"
    name_run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    # Right cell: Photo placeholder
    right_cell = header_table.rows[0].cells[1]
    right_cell.width = Inches(1.5)
    photo_para = right_cell.paragraphs[0]
    photo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if photo_path:
        try:
            photo_para.add_run().add_picture(photo_path, width=Cm(3.5), height=Cm(4.5))
        except:
            photo_para.add_run("[Foto einfügen]").font.size = Pt(8)
    else:
        photo_para.add_run("[Foto]").font.size = Pt(8)
        photo_para.add_run("[Foto]").font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Remove table borders
    for row in header_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 0}, bottom={"sz": 0}, start={"sz": 0}, end={"sz": 0})

    doc.add_paragraph()  # Spacer after header

    # --- PARSE CV TEXT INTO SECTIONS ---
    lines = cv_text.split("\n")
    current_section = None

    for line in lines:
        line = line.strip()

        if not line or line == "---":
            continue

        # Detect section headers (all caps or **SECTION**)
        if line.startswith("**") and line.endswith("**") and line.isupper():
            # New section header
            current_section = line.replace("**", "").strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)  # 6pt before section
            para.paragraph_format.space_after = Pt(0)  # 0pt after section
            run = para.add_run(current_section)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
            continue

        # Two-column entries (e.g., "MM/YYYY – MM/YYYY | Content")
        if "|" in line:
            parts = line.split("|", 1)
            date_part = parts[0].strip()
            content_part = parts[1].strip() if len(parts) > 1 else ""

            # Create a 2-column table for this entry
            entry_table = doc.add_table(rows=1, cols=2)
            entry_table.autofit = False

            # Left column: date
            left = entry_table.rows[0].cells[0]
            left.width = Inches(1.8)
            left_para = left.paragraphs[0]
            left_run = left_para.add_run(date_part)
            left_run.font.size = Pt(10)
            left_run.font.name = "Arial"

            # Right column: content
            right = entry_table.rows[0].cells[1]
            right.width = Inches(4.2)
            right_para = right.paragraphs[0]

            # Handle bold within content (**text**)
            content_parts = content_part.split("**")
            for i, part in enumerate(content_parts):
                if part:
                    run = right_para.add_run(part)
                    run.bold = (i % 2 == 1)
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

            # Remove table borders
            for row in entry_table.rows:
                for cell in row.cells:
                    set_cell_border(cell, top={"sz": 0}, bottom={"sz": 0}, start={"sz": 0}, end={"sz": 0})
                    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

        # Bullet points
        elif line.startswith("•") or line.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Cm(2.5)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(line.lstrip("•* "))
            run.font.size = Pt(10)
            run.font.name = "Arial"

        # Regular text line
        else:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)

            # Handle inline bold
            parts = line.split("**")
            for i, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    run.bold = (i % 2 == 1)
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

    # Save to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_cover_letter_document(cl_text: str, full_name: str) -> io.BytesIO:
    """DIN 5008 compliant cover letter, strictly 1 page."""
    if not cl_text or not isinstance(cl_text, str):
        raise ValueError(f"cl_text must be a non-empty string, got: {type(cl_text)}")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)

    for line in cl_text.split("\n"):
        line = line.strip()
        if not line or line == "---":
            doc.add_paragraph()
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)

        # Detect subject line (bold, no **markers**)
        if "Bewerbung als" in line or "bewerbung als" in line.lower():
            para.paragraph_format.space_before = Pt(12)
            run = para.add_run(line.replace("**", ""))
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Arial"
            continue

        # Handle inline bold
        parts = line.split("**")
        for i, part in enumerate(parts):
            if part:
                run = para.add_run(part)
                run.bold = (i % 2 == 1)
                run.font.size = Pt(10.5)
                run.font.name = "Arial"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def set_cell_border(cell, **kwargs):
    """
    Helper to remove table borders (makes the two-column layout look clean).
    Usage: set_cell_border(cell, top={"sz": 0}, bottom={"sz": 0}, ...)
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)
